import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# Konfigurasi Halaman Website
st.set_page_config(page_title="Edunexus RPP Generator", page_icon="📝", layout="wide")

# Fungsi: Minta AI buatkan teks RPP
def generate_rpp_with_ai(api_key, prompt_text):
    try:
        genai.configure(api_key=api_key)
        # Menggunakan model flash yang cepat dan stabil
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt_text)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# Fungsi: Bungkus teks jadi file Word
def create_docx(content, teacher_name, subject, topic):
    doc = Document()
    doc.add_heading(f'MODUL AJAR: {subject}', 0)
    doc.add_paragraph(f'Topik: {topic}')
    doc.add_paragraph(f'Guru: {teacher_name}')
    doc.add_paragraph('---')
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("📝 Edunexus RPP Generator (AI Powered)")
    st.markdown("Buat Modul Ajar (RPP) Kurikulum Merdeka otomatis dan download format Word.")

    # --- SIDEBAR (Kunci Masuk) ---
    with st.sidebar:
        st.header("🔑 Kunci Akses")
        api_key = st.text_input("Tempel Google Gemini API Key disini:", type="password")
        st.info("Butuh API Key agar aplikasi bisa menulis.")
        st.markdown("[Klik disini untuk dapat API Key Gratis](https://aistudio.google.com/app/apikey)")

    # --- INPUT DATA GURU ---
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("1. Data Sekolah")
        nama_guru = st.text_input("Nama Guru", "Guru Edunexus")
        sekolah = st.text_input("Nama Sekolah/Kursus", "Edunexus English Course")
        mapel = st.text_input("Mata Pelajaran", "Bahasa Inggris")
        kelas = st.text_input("Fase / Kelas", "Fase D / Kelas 7")

    with col2:
        st.subheader("2. Detail Materi")
        materi = st.text_input("Topik Materi", "Descriptive Text")
        durasi = st.text_input("Alokasi Waktu", "2 x 40 Menit")
        model = st.selectbox("Model Pembelajaran", ["Problem Based Learning", "Project Based Learning", "Discovery Learning"])
        profil = st.multiselect("Profil Pelajar Pancasila", 
                                ["Beriman & Bertaqwa", "Berkebinekaan Global", "Gotong Royong", "Mandiri", "Bernalar Kritis", "Kreatif"],
                                default=["Mandiri", "Kreatif"])

    st.markdown("---")

    # --- TOMBOL EKSEKUSI ---
    if st.button("✨ Buat RPP Sekarang", type="primary"):
        if not api_key:
            st.error("⚠️ Masukkan API Key dulu di menu sebelah kiri ya!")
        else:
            with st.spinner('Sedang mengetik RPP untuk Anda... (Mohon tunggu 20 detik)'):
                # Merangkai instruksi untuk AI
                profil_str = ", ".join(profil)
                prompt = f"""
                Buatkan Modul Ajar (RPP) Kurikulum Merdeka yang lengkap untuk:
                Guru: {nama_guru} di {sekolah}
                Mapel: {mapel}, Kelas: {kelas}
                Materi: {materi}
                Model: {model}
                Waktu: {durasi}
                Profil Pancasila: {profil_str}

                Struktur Wajib (Gunakan Bahasa Indonesia formal):
                1. TUJUAN PEMBELAJARAN
                2. PEMAHAMAN BERMAKNA & PERTANYAAN PEMANTIK
                3. KEGIATAN PEMBELAJARAN (Buka, Inti sesuai sintaks {model}, Penutup)
                4. ASESMEN (Formatif & Sumatif)
                5. LAMPIRAN (Ide Lembar Kerja Siswa singkat)

                Format output: Teks polos yang rapi (gunakan poin-poin), jangan gunakan tabel markdown agar mudah di-copy ke Word.
                """
                
                # Kirim ke AI
                hasil = generate_rpp_with_ai(api_key, prompt)

                if "Error" in hasil:
                    st.error(hasil)
                else:
                    st.success("Selesai! Silakan download file di bawah.")
                    
                    # Siapkan file Word
                    file_docx = create_docx(hasil, nama_guru, mapel, materi)
                    
                    # Tombol Download
                    st.download_button(
                        label="📥 Download File Word (.docx)",
                        data=file_docx,
                        file_name=f"RPP_{mapel}_{materi}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Tampilkan Preview
                    with st.expander("Lihat Bacaan RPP disini"):
                        st.write(hasil)

if __name__ == "__main__":
    main()